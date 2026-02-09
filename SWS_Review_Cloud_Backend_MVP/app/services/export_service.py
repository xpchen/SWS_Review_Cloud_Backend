import io
import json
from .. import db
from ..settings import settings
from ..core.deps import get_project_id_by_version_id

_schema = settings.DB_SCHEMA


def get_issues_for_export(version_id: int, status: str | None = None, severity: str | None = None) -> list[dict]:
    where = ["version_id = %(version_id)s"]
    params = {"version_id": version_id}
    if status:
        where.append("status = %(status)s")
        params["status"] = status
    if severity:
        where.append("severity = %(severity)s")
        params["severity"] = severity
    sql = f"""
    SELECT id, issue_type, severity, title, description, suggestion, confidence, status, page_no, created_at
    FROM {_schema}.review_issue
    WHERE {' AND '.join(where)}
    ORDER BY id DESC
    """
    return db.fetch_all(sql, params)


def get_issues_with_quotes_for_docx(version_id: int, status: str | None = None, severity: str | None = None) -> list[dict]:
    """获取带 evidence_quotes 的问题列表，用于 Word 问题清单导出。"""
    where = ["version_id = %(version_id)s"]
    params = {"version_id": version_id}
    if status:
        where.append("status = %(status)s")
        params["status"] = status
    if severity:
        where.append("severity = %(severity)s")
        params["severity"] = severity
    sql = f"""
    SELECT id, issue_type, severity, title, description, suggestion, confidence, status, page_no, evidence_quotes, created_at
    FROM {_schema}.review_issue
    WHERE {' AND '.join(where)}
    ORDER BY page_no ASC NULLS LAST, id ASC
    """
    rows = db.fetch_all(sql, params)
    for r in rows:
        eq = r.get("evidence_quotes")
        if isinstance(eq, str):
            try:
                r["evidence_quotes"] = json.loads(eq) if eq else []
            except Exception:
                r["evidence_quotes"] = [eq] if eq else []
        elif eq is None:
            r["evidence_quotes"] = []
    return rows


def get_document_title_for_version(version_id: int) -> str:
    """根据 version_id 获取文档标题。"""
    sql = f"""
    SELECT d.title
    FROM {_schema}.document d
    JOIN {_schema}.document_version dv ON dv.document_id = d.id
    WHERE dv.id = %(version_id)s
    """
    row = db.fetch_one(sql, {"version_id": version_id})
    return (row.get("title") or "文档").strip()


def get_outline_with_page_for_version(version_id: int) -> list[dict]:
    """获取版本大纲节点（含 page_no），按 order_index 排序。"""
    sql = f"""
    SELECT n.id, n.node_no, n.title, n.level, n.parent_id, n.order_index,
           COALESCE(
             (SELECT bpa.page_no
              FROM {_schema}.doc_block b
              JOIN {_schema}.block_page_anchor bpa ON bpa.block_id = b.id
              WHERE b.outline_node_id = n.id AND b.block_type = 'HEADING'
              ORDER BY bpa.page_no ASC
              LIMIT 1),
             1
           ) AS page_no
    FROM {_schema}.doc_outline_node n
    WHERE n.version_id = %(version_id)s
    ORDER BY n.order_index ASC
    """
    return db.fetch_all(sql, {"version_id": version_id})


def _issue_type_to_review_type_label(issue_type: str) -> str:
    """将 issue_type 映射为 Word 中的审查类型标签。"""
    if not issue_type:
        return "其他审查"
    t = issue_type.upper()
    if "CONSISTENCY" in t:
        return "一致性审查"
    if "BUSINESS_LOGIC" in t:
        return "业务逻辑审查"
    if "FORMAT" in t:
        return "格式审查"
    if "CONTENT" in t:
        return "内容审查"
    if "SUM_MISMATCH" in t or "FORMULA" in t or "PERCENTAGE" in t or "UNIT_INCONSISTENT" in t or "KEY_FIELD" in t:
        return "表内计算审查"
    return "其他审查"


def _is_formal_review(issue_type: str) -> bool:
    """形式审查：格式、内容等。"""
    if not issue_type:
        return False
    t = issue_type.upper()
    return "FORMAT" in t or "CONTENT" in t


def _is_tech_review(issue_type: str) -> bool:
    """技术审查：一致性、业务逻辑、表内计算等。"""
    if not issue_type:
        return False
    t = issue_type.upper()
    return "CONSISTENCY" in t or "BUSINESS_LOGIC" in t or "SUM_MISMATCH" in t or "FORMULA" in t or "PERCENTAGE" in t or "UNIT_INCONSISTENT" in t or "KEY_FIELD" in t or "MISSING_SECTION" in t


def _outline_node_path(node: dict, node_by_id: dict) -> str:
    """求单个大纲节点从根到该节点的路径字符串（多行）。"""
    path_parts = []
    current = node
    while current:
        part = ((current.get("node_no") or "") + " " + (current.get("title") or "").strip()).strip()
        if part:
            path_parts.append(part)
        pid = current.get("parent_id")
        current = node_by_id.get(pid) if pid else None
    path_parts.reverse()
    return "\n".join(path_parts)


def _assign_issue_to_outline_section(issue_page: int | None, outline_nodes: list[dict]) -> str:
    """根据问题页码归属到大纲节点，返回该节点的完整路径字符串（如 '1. 项目概况 \\n1.1 项目组成'）。"""
    if not outline_nodes:
        return ""
    node_by_id = {n["id"]: n for n in outline_nodes}
    if issue_page is None:
        first = outline_nodes[0]
        return _outline_node_path(first, node_by_id)
    # 找到 page_no <= issue_page 的最后一个节点
    chosen = None
    for n in outline_nodes:
        p = n.get("page_no")
        if p is not None:
            try:
                if int(p) <= issue_page:
                    chosen = n
            except (TypeError, ValueError):
                pass
    if chosen is None:
        chosen = outline_nodes[0]
    return _outline_node_path(chosen, node_by_id)


def build_issues_docx(version_id: int, status: str | None = None, severity: str | None = None) -> bytes:
    """生成「问题清单」格式的 Word 文档（.docx）。"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc_title = get_document_title_for_version(version_id)
    outline_nodes = get_outline_with_page_for_version(version_id)
    issues = get_issues_with_quotes_for_docx(version_id, status=status, severity=severity)

    doc = Document()
    # 正文样式：宋体 小四
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题：{文档标题}问题清单
    title_text = (doc_title if doc_title.endswith("问题清单") else doc_title + "问题清单").strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph()

    # 审查错误统计表
    formal_count = sum(1 for i in issues if _is_formal_review(i.get("issue_type")))
    tech_count = sum(1 for i in issues if _is_tech_review(i.get("issue_type")))
    total_count = len(issues)
    doc.add_paragraph("审查错误统计表")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "审查类型"
    table.rows[0].cells[1].text = "错误"
    table.rows[1].cells[0].text = "技术审查"
    table.rows[1].cells[1].text = str(tech_count)
    table.rows[2].cells[0].text = "形式审查"
    table.rows[2].cells[1].text = str(formal_count)
    table.rows[3].cells[0].text = "总计"
    table.rows[3].cells[1].text = str(total_count)
    doc.add_paragraph()

    # 按大纲章节分组问题
    section_to_issues: dict[str, list[dict]] = {}
    for issue in issues:
        page_no = issue.get("page_no")
        if page_no is not None:
            try:
                page_no = int(page_no)
            except (TypeError, ValueError):
                page_no = 1
        section_path = _assign_issue_to_outline_section(page_no, outline_nodes)
        if not section_path.strip():
            section_path = "其他"
        if section_path not in section_to_issues:
            section_to_issues[section_path] = []
        section_to_issues[section_path].append(issue)

    # 按大纲顺序排序章节（用 outline 的 order 决定顺序，未匹配的放最后）
    node_by_id = {n["id"]: n for n in outline_nodes}
    order_keys = []
    for n in outline_nodes:
        path = _outline_node_path(n, node_by_id)
        if path and path not in order_keys:
            order_keys.append(path)
    for sec in section_to_issues:
        if sec not in order_keys:
            order_keys.append(sec)
    ordered_sections = [s for s in order_keys if s in section_to_issues]

    for section_path in ordered_sections:
        issue_list = section_to_issues[section_path]
        # 章节标题（多行时用换行）
        for line in section_path.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        for issue in issue_list:
            review_label = _issue_type_to_review_type_label(issue.get("issue_type"))
            title = (issue.get("title") or "").strip()
            description = (issue.get("description") or "").strip()
            quotes = issue.get("evidence_quotes") or []
            quote_lines = []
            for q in quotes:
                if isinstance(q, dict):
                    quote_lines.append((q.get("text") or q.get("quote") or str(q)).strip())
                else:
                    quote_lines.append(str(q).strip())
            quote_text = "\n".join(q for q in quote_lines if q)

            doc.add_paragraph(review_label)
            doc.add_paragraph(title)
            if quote_text:
                doc.add_paragraph("原文片段：" + quote_text)
            if description:
                doc.add_paragraph("推理过程: " + description)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_issues_xlsx(version_id: int, status: str | None = None, severity: str | None = None) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, Alignment
    rows = get_issues_for_export(version_id, status, severity)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "审查问题"
    headers = ["ID", "类型", "严重程度", "标题", "描述", "建议", "置信度", "状态", "页码", "创建时间"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
        ws.cell(row=1, column=col).font = Font(bold=True)
    for row_idx, r in enumerate(rows, 2):
        ws.cell(row=row_idx, column=1, value=r.get("id"))
        ws.cell(row=row_idx, column=2, value=r.get("issue_type"))
        ws.cell(row=row_idx, column=3, value=r.get("severity"))
        ws.cell(row=row_idx, column=4, value=r.get("title"))
        ws.cell(row=row_idx, column=5, value=r.get("description"))
        ws.cell(row=row_idx, column=6, value=r.get("suggestion"))
        ws.cell(row=row_idx, column=7, value=r.get("confidence"))
        ws.cell(row=row_idx, column=8, value=r.get("status"))
        ws.cell(row=row_idx, column=9, value=r.get("page_no"))
        ws.cell(row=row_idx, column=10, value=str(r.get("created_at")) if r.get("created_at") else "")
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


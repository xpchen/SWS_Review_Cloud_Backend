import hashlib
import io
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from .. import db
from ..settings import settings
from ..services.file_service import get_file_object
from ..services import kb_service
from ..storage import get_storage
from .app import app

_schema = settings.DB_SCHEMA


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 100, page_boundaries: list[tuple[int, int, int]] | None = None) -> list[tuple[str, dict]]:
    """
    切分文本为chunks，返回(chunk_text, meta)列表。
    meta包含page_start, page_end等信息。
    
    Args:
        text: 要切分的文本
        chunk_size: chunk大小
        overlap: 重叠长度
        page_boundaries: 页边界列表 [(char_start, char_end, page_no), ...]，用于PDF分页映射
    """
    import bisect
    
    parts = []
    start = 0
    chunk_idx = 0
    
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk_text = text[start:end]
        meta = {
            "chunk_index": chunk_idx,
            "char_start": start,
            "char_end": end,
        }
        
        # 使用页边界+二分查找确定page_start和page_end
        if page_boundaries:
            # 查找start所在的页
            page_start = None
            page_end = None
            
            # 二分查找：找到start所在的页
            idx = bisect.bisect_right([b[0] for b in page_boundaries], start) - 1
            if idx >= 0 and idx < len(page_boundaries):
                boundary_start, boundary_end, page_no = page_boundaries[idx]
                if boundary_start <= start < boundary_end:
                    page_start = page_no
            
            # 查找end所在的页
            idx = bisect.bisect_right([b[0] for b in page_boundaries], end) - 1
            if idx >= 0 and idx < len(page_boundaries):
                boundary_start, boundary_end, page_no = page_boundaries[idx]
                if boundary_start <= end < boundary_end:
                    page_end = page_no
            
            if page_start:
                meta["page_start"] = page_start
            if page_end:
                meta["page_end"] = page_end
        
        parts.append((chunk_text, meta))
        start = end - overlap if end < len(text) else len(text)
        chunk_idx += 1
    
    return parts


def _extract_text_from_pdf(raw: bytes, filename: str) -> tuple[str, list[tuple[int, int, int]]]:
    """
    从PDF提取文本，返回(完整文本, 页边界列表)。
    页边界列表用于二分查找：[(char_start, char_end, page_no), ...]
    
    Returns:
        (full_text, page_boundaries)
        page_boundaries: [(char_start, char_end, page_no), ...]
    """
    doc_pdf = fitz.open(stream=raw, filetype="pdf")
    full_text_parts = []
    page_boundaries = []
    
    char_offset = 0
    
    for page_no in range(len(doc_pdf)):
        page = doc_pdf[page_no]
        page_text = page.get_text()
        
        # 记录页边界
        page_start = char_offset
        page_end = char_offset + len(page_text)
        page_boundaries.append((page_start, page_end, page_no + 1))
        
        full_text_parts.append(page_text)
        char_offset = page_end + 2  # +2 for "\n\n"
    
    doc_pdf.close()
    full_text = "\n\n".join(full_text_parts)
    return full_text, page_boundaries


def _extract_text_from_docx(raw: bytes, filename: str) -> tuple[str, None]:
    """
    从DOCX提取文本，返回(完整文本, 段落信息列表)。
    """
    docx_doc = DocxDocument(io.BytesIO(raw))
    text_parts = []
    
    for para in docx_doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            text_parts.append(text)
    
    # 也提取表格文本
    for table in docx_doc.tables:
        table_texts = []
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                table_texts.append(" | ".join(row_texts))
        if table_texts:
            text_parts.append("\n".join(table_texts))
    
    full_text = "\n\n".join(text_parts)
    return full_text, []


@app.task(bind=True)
def index_kb_source_task(self, source_id: int):
    src = kb_service.get_kb_source(source_id)
    if not src or src["status"] != "PROCESSING":
        return
    fo = get_file_object(src["file_id"])
    if not fo:
        kb_service.set_kb_source_failed(source_id, "File not found")
        return
    
    storage = get_storage()
    try:
        stream = storage.get_object(fo["object_key"])
        if not stream:
            kb_service.set_kb_source_failed(source_id, "Object not found")
            return
        raw = stream.read()
        if hasattr(stream, "close"):
            stream.close()
    except Exception as e:
        kb_service.set_kb_source_failed(source_id, str(e))
        return
    
    content_type = fo.get("content_type") or ""
    filename = fo.get("filename") or ""
    
    # 根据content_type选择提取方法
    text = ""
    page_infos = []
    
    page_boundaries = None
    
    if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
        try:
            text, page_boundaries = _extract_text_from_pdf(raw, filename)
        except Exception as e:
            kb_service.set_kb_source_failed(source_id, f"PDF extraction failed: {str(e)}")
            return
    elif content_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ] or filename.lower().endswith((".docx", ".doc")):
        try:
            text, _ = _extract_text_from_docx(raw, filename)
            # DOCX没有分页信息
            page_boundaries = None
        except Exception as e:
            kb_service.set_kb_source_failed(source_id, f"DOCX extraction failed: {str(e)}")
            return
    else:
        # 尝试作为纯文本解码
        try:
            text = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
            page_boundaries = None
        except Exception as e:
            kb_service.set_kb_source_failed(source_id, f"Text extraction failed: {str(e)}")
            return
    
    if not text.strip():
        kb_service.set_kb_source_ready(source_id)
        return
    
    # 切分chunks（传入page_boundaries用于二分查找）
    chunks_with_meta = _chunk_text(text, page_boundaries=page_boundaries)
    
    # 入库chunks
    for chunk_text, meta in chunks_with_meta:
        meta["doc"] = filename
        meta["clause_hint"] = ""  # 可以后续用AI提取
        
        h = hashlib.sha256(chunk_text.encode("utf-8")).hexdigest()
        kb_service.insert_chunk(source_id, chunk_text, meta, h)
    
    kb_service.set_kb_source_ready(source_id)
